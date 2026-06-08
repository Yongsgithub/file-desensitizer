#!/usr/bin/env python3
"""
Word 文档（.docx）脱敏处理器

支持对 .docx 文件中的段落文本和表格内容进行脱敏处理。
处理方式：
1. 遍历所有段落和表格单元格
2. 对文本内容进行脱敏
3. 生成脱敏后的新 .docx 文件
"""

import sys
import os
import json
import re
from pathlib import Path
from typing import Dict, Any, List

from docx import Document
from docx.shared import RGBColor, Pt
from docx.oxml.ns import qn

sys.path.insert(0, str(Path(__file__).parent))
from desensitizer import TextDesensitizer


def process_paragraph(paragraph, records: List[Dict]):
    """
    处理单个段落中的文本。

    先将段落内所有 run 的文本拼接为完整文本，再进行脱敏替换。
    这样可以处理跨 run 分布的敏感信息（如手机号、学号被 Word
    拆分成多个 run 片段的情况）。

    脱敏后将文本按 run 边界切分回各 run，并用红色高亮标记脱敏文本。

    注意：当前所有脱敏模式均保持字符长度不变（身份证18→18、
    手机号11→11、学号11→11、邮编6→6等），因此按原始 run 长度
    切分不会产生偏移。
    """
    runs = paragraph.runs
    if not runs:
        return

    # 收集所有 run 的文本和长度
    run_texts = []
    for run in runs:
        run_texts.append(run.text)

    full_text = ''.join(run_texts)
    if not full_text.strip():
        return

    # 对全文进行脱敏
    desensitized, run_records = TextDesensitizer.desensitize_text(full_text)
    if desensitized == full_text:
        return

    records.extend(run_records)

    if len(desensitized) == len(full_text):
        # 快速路径：长度不变（手机号/学号/身份证/邮编/邮箱/出生年月/姓名
        # 等所有数字型脱敏均保持长度不变），按原始 run 边界切分即可保持
        # 各 run 的格式属性（字体、大小、加粗等）。
        pos = 0
        for i, run in enumerate(runs):
            run_len = len(run_texts[i])
            if run_len > 0:
                new_text = desensitized[pos:pos + run_len]
                if new_text != run_texts[i]:
                    run.text = new_text
                    run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
            pos += run_len
    else:
        # 长度变化路径（地址脱敏等会导致文本缩短）：
        # 将所有 run 清空，用第一个 run 承载全文并标红。
        # 这会丢失原有格式化，但保证脱敏结果正确。
        for run in runs:
            run.text = ''
        if runs:
            runs[0].text = desensitized
            runs[0].font.color.rgb = RGBColor(0xFF, 0x00, 0x00)


def process_table(table, records: List[Dict]):
    """处理表格中的所有单元格"""
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                process_paragraph(paragraph, records)


def process_docx(
    input_path: str,
    output_dir: str = None,
    mark_color: bool = True
) -> Dict[str, Any]:
    """
    处理 Word 文档脱敏。

    Args:
        input_path: 输入 .docx 文件路径
        output_dir: 输出目录（默认为输入文件同目录）
        mark_color: 是否用红色标记脱敏文本

    Returns:
        处理结果字典
    """
    input_path = str(Path(input_path).resolve())
    input_file = Path(input_path)

    if not input_file.exists():
        return {"success": False, "error": f"文件不存在: {input_path}"}

    if not input_file.suffix.lower() in ('.docx',):
        return {"success": False, "error": f"不支持的文件格式: {input_file.suffix}，仅支持 .docx"}

    if output_dir is None:
        output_dir = input_file.parent
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    result = {
        "success": True,
        "input_path": input_path,
        "output_path": None,
        "text_output_path": None,
        "records": [],
    }

    try:
        doc = Document(input_path)
    except Exception as e:
        return {"success": False, "error": f"无法打开文档: {e}"}

    all_records = []

    # 1. 处理所有段落（正文、页眉页脚等）
    for paragraph in doc.paragraphs:
        process_paragraph(paragraph, all_records)

    # 2. 处理表格
    for table in doc.tables:
        process_table(table, all_records)

    # 3. 处理页眉
    for section in doc.sections:
        try:
            header = section.header
            if header:
                for paragraph in header.paragraphs:
                    process_paragraph(paragraph, all_records)
                for table in header.tables:
                    process_table(table, all_records)
        except Exception:
            pass

    # 4. 处理页脚
    for section in doc.sections:
        try:
            footer = section.footer
            if footer:
                for paragraph in footer.paragraphs:
                    process_paragraph(paragraph, all_records)
                for table in footer.tables:
                    process_table(table, all_records)
        except Exception:
            pass

    # 去重记录
    seen = set()
    unique_records = []
    for r in all_records:
        key = (r['category'], r['original'])
        if key not in seen:
            seen.add(key)
            unique_records.append(r)

    result["records"] = unique_records

    # 5. 保存脱敏后的文档
    output_path = output_dir / f"{input_file.stem}_desensitized.docx"
    doc.save(str(output_path))
    result["output_path"] = str(output_path)

    # 6. 保存脱敏记录文本
    text_output_path = output_dir / f"{input_file.stem}_desensitized_record.txt"
    with open(text_output_path, 'w', encoding='utf-8') as f:
        f.write(f"=== Word 文档脱敏记录 ===\n")
        f.write(f"原文件: {input_path}\n")
        f.write(f"脱敏文件: {output_path}\n\n")
        for r in unique_records:
            f.write(f"[{r['category']}] {r['original']} → {r['masked']}\n")
    result["text_output_path"] = str(text_output_path)

    return result


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("用法: python docx_processor.py <docx文件路径> [输出目录]")
        sys.exit(1)

    input_path = sys.argv[1]
    output_dir = sys.argv[2] if len(sys.argv) > 2 else None

    result = process_docx(input_path, output_dir)
    print(json.dumps(result, ensure_ascii=False, indent=2))
